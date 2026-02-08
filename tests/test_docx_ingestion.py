import io
from app import app
from storage import get_connection
from docx import Document


def count_fragments():
    conn = get_connection()
    try:
        return conn.execute("SELECT COUNT(*) FROM fragments").fetchone()[0]
    finally:
        conn.close()


def test_docx_ingestion_creates_paragraph_fragments():
    client = app.test_client()

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(buf)
    buf.seek(0)

    before = count_fragments()

    client.post(
        "/disassembler",
        data={"files": [(buf, "test.docx")]},
        content_type="multipart/form-data",
    )

    after = count_fragments()

    assert after == before + 2
